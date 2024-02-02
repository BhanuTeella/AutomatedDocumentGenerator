import random
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch


def generate_question():
    # Generate 8 to 10 odd numbers between 1 to 9
    odd_numbers = random.sample([n for n in range(1, 10) if n % 2 == 1]*3, random.randint(8, 10))
    
    # Calculate the sum of odd numbers
    answer = sum(odd_numbers)
    
    # Generate remaining numbers
    even_numbers = [random.randrange(2, 10, 2) for _ in range(50 - len(odd_numbers))]
    
    # Combine odd and even numbers
    numbers = odd_numbers + even_numbers
    random.shuffle(numbers)
        
    # Generate 4 options, one of which is the correct sum
    options = [answer] + [answer+random.choice([n for n in range(-9, 0)] + [n for n in range(1, 10)]) for _ in range(3)]
    random.shuffle(options)
    
    # Return the question and options as a tuple
    return (numbers, options, answer)



import docx

# Create a new Word document
doc = docx.Document()

# Access the section properties
section = doc.sections[0]
section.left_margin = docx.shared.Inches(0.75)
section.right_margin = docx.shared.Inches(0.5)
section.top_margin = docx.shared.Inches(0.75)
section.bottom_margin = docx.shared.Inches(0.75)

# Modify the font properties
font = doc.styles['Normal'].font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(12)

for set in range(30):
    correct_answer_list=[]
    for i in range(30):
        question, options, answer = generate_question()
        correct_answer_list.append(answer)
        # Add the question and options to the same paragraph
        p = doc.add_paragraph()
        p.add_run(f"{i+1}:\t {' '.join(str(num) for num in question)}")
        p.add_run("\n")
        p.add_run("\t".join([f"\t{j+1}. {option}" for j, option in enumerate(options)]))
    doc.add_paragraph("Correct Answers:")
    doc.add_paragraph("          ".join([f"{j+1}. {option}" for j, option in enumerate(correct_answer_list)]))

# Save the document
doc.save("final.docx")



# Loop through the questions and options
'''for i in range(30):
    question, options = generate_question()
    
    # Add the question to the document
    doc.add_paragraph(f"{i+1}:\t {' '.join(str(num) for num in question)}")
    
    # Add the options to the document
    option_text = "\t".join([f"{j+1}. {option}" for j, option in enumerate(options)])
    doc.add_paragraph(option_text)
'''
#doc.add_heading("Document Heading", level=0)



'''# Generate 30 questions
for i in range(30):
    question, options = generate_question()
    print(f"Question {i+1}: {' '.join(str(num) for num in question)}")
    
    for j, option in enumerate(options):
        print(f"{j+1}. {option}", end="\t")
    
    print() # add a blank line'''